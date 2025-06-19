# cover_letter_generator.py
from docx import Document
from fastapi.responses import FileResponse
import tempfile

def generate_cover_letter(data):
    name = data.get("name")
    email = data.get("email")
    phone = data.get("phone")
    company = data.get("company")
    job_title = data.get("job_title")
    resume_summary = data.get("summary", "")
    tone = data.get("tone", "formal")

    doc = Document()

    doc.add_paragraph(f"{name}\n{email} | {phone}\n")
    doc.add_paragraph(f"To,\nHiring Manager\n{company}\n")
    doc.add_paragraph("\n")

    doc.add_paragraph(f"Dear Hiring Manager,")

    intro = f"I am writing to express my interest in the {job_title} position at {company}. " \
           f"With a background in {resume_summary}, I am confident in my ability to contribute effectively to your team."
    doc.add_paragraph(intro)

    body = (
        "I have developed strong skills through my academic and project experiences that align well with the requirements of this role. "
        "My ability to quickly learn and apply new technologies, along with my passion for continuous improvement, makes me a strong candidate."
    )
    doc.add_paragraph(body)

    closing = (
        "I would welcome the opportunity to further discuss how I can contribute to your team. "
        "Thank you for considering my application. I look forward to hearing from you."
    )
    doc.add_paragraph(closing)

    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(name)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(
            tmp.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{name.replace(' ', '_')}_Cover_Letter.docx"
        )
