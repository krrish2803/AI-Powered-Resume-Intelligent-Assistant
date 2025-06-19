from docx import Document
from fastapi.responses import FileResponse
import tempfile

def generate_resume(data):
    name = data.get("name")
    email = data.get("email")
    phone = data.get("phone")
    experience = data.get("experience",[])
    education = data.get("education",[])
    skills = data.get("skills",[])
    projects = data.get("projects",[])

    document = Document()

    document.add_heading(name, 0)
    document.add_paragraph(f"Email: {email} | Phone: {phone}")

    document.add_heading("Education", level=1)
    for edu in education:
        degree = edu.get("degree")
        institution = edu.get("institution")
        graduation_year = edu.get("graduation_year")
        document.add_paragraph(f"{degree} in {institution}, {graduation_year}")

    document.add_heading("Skills",level=1)
    document.add_paragraph(",".join(skills))

    if experience:
        document.add_heading("Experience",level=1)
        for exp in experience:
            title = exp.get("title")
            company = exp.get("company")
            duration = exp.get("duration")

            document.add_paragraph(f"{title} at {company} ({duration})")

    if projects:
        document.add_heading("Projects",level=1)
        for project in projects:
            title = project.get("title")
            description = project.get("description")

            document.add_paragraph(f"{title}")
            document.add_paragraph(f"{description}")

        # Save to temporary file and return
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        document.save(tmp.name)
        return FileResponse(
            tmp.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{name.replace(' ', '_')}_Resume.docx"
        )        

